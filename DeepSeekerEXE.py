# @Author        : jepyh
# @Time          : ${2025-02-13} ${16:20}
# @Version       : 1.0

import tkinter as tk
from tkinter import filedialog
import tkinter.messagebox
from docx import Document
from openai import OpenAI
import json

def call_deepseek(call_text):
# for backward compatibility, you can still use `https://api.deepseek.com/v1` as `base_url`.
    client = OpenAI(api_key="输入您的api-key", base_url="https://api.deepseek.com")

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": call_text},
    ],
        max_tokens=1024,
        temperature=0.7,
        stream=False
    )
    #print(response.choices[0].message.content)
    return response.choices[0].message.content

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx *.doc")])
    if file_path:
        flag_success=append_text_to_file(file_path)
    if flag_success:
        tkinter.messagebox.showinfo('提示','扩写成功')

def append_text_to_file(file_path):
    try:
        doc = Document(file_path)
        full_text=''
        paragraphs=doc.paragraphs
        for p in paragraphs:
            print(p.text)
            full_text+=str(p.text)
        #print(full_text)#文档文本
        doc.add_paragraph(call_deepseek(full_text))
        doc.save(file_path)
        return True
    except PermissionError:
        tkinter.messagebox.showinfo('提示','请保存、关闭文档后重试')
        print('请保存、关闭文档后重试')
        return False
    except json.decoder.JSONDecodeError as e:
        tkinter.messagebox.showinfo('提示','服务器繁忙，请稍后再试')
        print('服务器繁忙，请稍后再试')
        return False

root = tk.Tk()
root.title("Word文档扩写工具DeepSeek")
width = root.winfo_screenwidth()
height = root.winfo_screenheight()
root.geometry("%dx%d+%d+%d" % (int(350), int(80), int(50), int(50)))

select_button = tk.Button(root, text="请选择需要扩写的文档文件", command=select_file)
select_button.pack(pady=20)

root.mainloop()